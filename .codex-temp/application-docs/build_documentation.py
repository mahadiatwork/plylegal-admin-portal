from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\Work Projects\validify-pro-admin-portal")
OUTPUT = ROOT / "output" / "ValidifyPro_Admin_Portal_Application_Documentation.docx"

NAVY = "17372E"
GREEN = "4F726B"
GREEN_LIGHT = "EAF2EE"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "68747F"
DARK = "1F2933"
BORDER = "D6DEE8"
WHITE = "FFFFFF"
AMBER = "7A5A00"
AMBER_LIGHT = "FFF8E1"
RED = "9B1C1C"
RED_LIGHT = "FDECEC"
CODE_FILL = "F5F7FA"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    set_table_borders(table)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeatable_style(style, font="Calibri", size=11, color=DARK, bold=False):
    style.font.name = font
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_style_spacing(style, before=0, after=0, line=1.0, keep_with_next=False):
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MID_GRAY)


def add_numbering_definition(document, kind="bullet"):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def configure_document(document):
    document.core_properties.title = "ValidifyPro Admin Portal - Application and Technical Documentation"
    document.core_properties.subject = "Features, architecture, APIs, data storage, configuration, and repository structure"
    document.core_properties.author = "ValidifyPro Engineering Documentation"
    document.core_properties.keywords = "ValidifyPro, Next.js, Firebase, Zoho CRM, WorkDrive, admin portal"

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    set_repeatable_style(normal, size=11, color=DARK)
    set_style_spacing(normal, before=0, after=6, line=1.25)

    title = styles["Title"]
    set_repeatable_style(title, size=30, color=NAVY, bold=True)
    set_style_spacing(title, before=0, after=8, line=1.0, keep_with_next=True)

    subtitle = styles["Subtitle"]
    set_repeatable_style(subtitle, size=14, color=GREEN)
    set_style_spacing(subtitle, before=0, after=14, line=1.1, keep_with_next=True)

    h1 = styles["Heading 1"]
    set_repeatable_style(h1, size=16, color=BLUE, bold=True)
    set_style_spacing(h1, before=18, after=10, line=1.0, keep_with_next=True)

    h2 = styles["Heading 2"]
    set_repeatable_style(h2, size=13, color=BLUE, bold=True)
    set_style_spacing(h2, before=14, after=7, line=1.0, keep_with_next=True)

    h3 = styles["Heading 3"]
    set_repeatable_style(h3, size=12, color=BLUE_DARK, bold=True)
    set_style_spacing(h3, before=10, after=5, line=1.0, keep_with_next=True)

    caption = styles["Caption"]
    set_repeatable_style(caption, size=9, color=MID_GRAY, bold=True)
    set_style_spacing(caption, before=4, after=4, line=1.0, keep_with_next=True)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    set_repeatable_style(code, font="Consolas", size=8.5, color=DARK)
    set_style_spacing(code, before=3, after=6, line=1.0)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.08)

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", 1)
    else:
        table_text = styles["Table Text"]
    set_repeatable_style(table_text, size=9.2, color=DARK)
    set_style_spacing(table_text, before=0, after=2, line=1.1)

    header = section.header
    p = header.paragraphs[0]
    p.text = "VALIDIFYPRO ADMIN PORTAL  |  APPLICATION DOCUMENTATION"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        set_run_font(run, size=8, color=MID_GRAY, bold=True)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("Page ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    add_field(p, "PAGE")
    run = p.add_run(" of ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    add_field(p, "NUMPAGES")

    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def add_text(document, text="", *, bold=False, italic=False, color=None, size=None, style=None, align=None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        color=color,
        bold=bold,
        italic=italic,
    )
    return paragraph


def add_mixed_paragraph(document, segments, style=None, after=None, keep_with_next=None):
    paragraph = document.add_paragraph(style=style)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if keep_with_next is not None:
        paragraph.paragraph_format.keep_with_next = keep_with_next
    for segment in segments:
        if isinstance(segment, str):
            segment = {"text": segment}
        run = paragraph.add_run(segment.get("text", ""))
        set_run_font(
            run,
            name=segment.get("font", "Calibri"),
            size=segment.get("size"),
            color=segment.get("color"),
            bold=segment.get("bold"),
            italic=segment.get("italic"),
        )
    return paragraph


def add_bullets(document, items):
    num_id = add_numbering_definition(document, "bullet")
    for item in items:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, num_id)
        if isinstance(item, tuple):
            label, body = item
            run = paragraph.add_run(f"{label}: ")
            set_run_font(run, bold=True, color=NAVY)
            run = paragraph.add_run(body)
            set_run_font(run)
        else:
            run = paragraph.add_run(item)
            set_run_font(run)


def add_numbered(document, items):
    num_id = add_numbering_definition(document, "decimal")
    for item in items:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, num_id)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_callout(document, label, body, kind="info"):
    palette = {
        "info": (GREEN_LIGHT, GREEN),
        "warning": (AMBER_LIGHT, AMBER),
        "risk": (RED_LIGHT, RED),
    }
    fill, accent = palette[kind]
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    shade_paragraph(paragraph, fill, accent)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, bold=True, color=accent)
    run = paragraph.add_run(body)
    set_run_font(run, color=DARK)
    return paragraph


def add_code_block(document, text):
    paragraph = document.add_paragraph(style="Code Block")
    paragraph.paragraph_format.keep_together = True
    shade_paragraph(paragraph, CODE_FILL, BORDER)
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=8.5, color=DARK)
    return paragraph


def add_table(document, headers, rows, widths_dxa):
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, BLUE_GRAY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.style = document.styles["Table Text"]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.2, color=NAVY, bold=True)

    for row in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(row):
            cell = cells[index]
            if len(table.rows) % 2 == 0:
                shade_cell(cell, "FAFBFC")
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["Table Text"]
            paragraph.paragraph_format.space_after = Pt(0)
            if isinstance(value, tuple):
                label, body = value
                run = paragraph.add_run(label)
                set_run_font(run, size=9.2, color=NAVY, bold=True)
                run = paragraph.add_run(body)
                set_run_font(run, size=9.2, color=DARK)
            else:
                run = paragraph.add_run(str(value))
                set_run_font(run, size=9.2, color=DARK)
    return table


def add_table_caption(document, text):
    paragraph = document.add_paragraph(style="Caption")
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MID_GRAY, bold=True)
    return paragraph


def add_page_break(document):
    document.add_page_break()


def add_section_intro(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=11.3, color=MID_GRAY, italic=True)


def build_document():
    document = Document()
    configure_document(document)

    # Cover
    add_text(
        document,
        "TECHNICAL REFERENCE",
        bold=True,
        color=GREEN,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    ).paragraph_format.space_before = Pt(104)
    add_text(
        document,
        "ValidifyPro Admin Portal",
        style="Title",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(
        document,
        "Application and Technical Documentation",
        style="Subtitle",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p = add_text(
        document,
        "Features | Architecture | APIs | Data Storage | Configuration | Operations",
        color=MID_GRAY,
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p.paragraph_format.space_after = Pt(88)
    add_text(
        document,
        "Repository snapshot",
        bold=True,
        color=NAVY,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(
        document,
        "main @ 905a184, including the inspected working-tree state",
        color=MID_GRAY,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text(
        document,
        "Prepared 23 July 2026",
        color=MID_GRAY,
        size=9.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_callout(
        document,
        "Scope",
        "This document is derived from the current repository implementation. Credential values, private keys, access tokens, and service-account contents are intentionally excluded.",
    )

    add_page_break(document)

    # Document map
    document.add_heading("Document map", level=1)
    add_section_intro(
        document,
        "Use this page as a fast index. The document distinguishes implemented behavior from installed-but-unused packages and from adjacent Firestore rules.",
    )
    directory = [
        ("1", "Application overview and architecture"),
        ("2", "Feature catalog and user journeys"),
        ("3", "Application pages and technical stack"),
        ("4", "Internal API reference"),
        ("5", "External APIs and integrations"),
        ("6", "Data storage and information ownership"),
        ("7", "Core data flows"),
        ("8", "Authentication, authorization, and security"),
        ("9", "Configuration and environment variables"),
        ("10", "Folder structure and source ownership"),
        ("11", "Development, deployment, and operations"),
        ("12", "Verification results, limitations, and risks"),
        ("Appendix", "Data shapes and source map"),
    ]
    add_table_caption(document, "Section directory")
    add_table(document, ["Section", "Contents"], directory, [1440, 7920])

    document.add_heading("Quick facts", level=2)
    quick_facts = [
        ("Application", "ValidifyPro Admin Portal / Matter Viewer Portal"),
        ("Framework", "Next.js 16.2.4 App Router with React 19.2.4"),
        ("Primary database", "Google Firebase Firestore through Firebase Admin SDK"),
        ("CRM", "Zoho CRM v7; CRM v8 is available for attachment upload"),
        ("File storage", "Zoho WorkDrive; Firestore stores resource metadata and public/share URLs"),
        ("Supported resource types", "Files, links, and notes; logical folders are available in visa templates"),
        ("Matter identifiers", "Firebase application document ID or Zoho Deal ID"),
        ("Visa experiences", "Partner, Protection, Skills in Demand (subclass 482), and Employer Nomination (subclass 186)"),
        ("UI language", "JavaScript and JSX, with server-side Node.js route handlers"),
        ("Current access posture", "Matter and admin route handlers are not effectively access-gated in the current source"),
    ]
    add_table_caption(document, "Application at a glance")
    add_table(document, ["Topic", "Current implementation"], quick_facts, [2520, 6840])

    add_page_break(document)

    # 1. Overview
    document.add_heading("1. Application overview and architecture", level=1)
    add_section_intro(
        document,
        "The portal gives staff a consolidated view of an immigration matter while keeping operational records in Firestore and file content in Zoho WorkDrive.",
    )
    document.add_heading("1.1 Purpose", level=2)
    add_text(
        document,
        "The application is a Next.js-based matter review and resource administration portal. A staff member enters either a Firebase Application ID or a Zoho Deal ID, then reviews questionnaire data, manages matter-specific and shared resources, uploads a prepared PDF for document review, and tracks correction status.",
    )
    add_bullets(
        document,
        [
            ("Matter review", "Resolve one matter, show application metadata, and calculate questionnaire completion."),
            ("Questionnaire review", "Present stored answers as read-only fields with profile-aware navigation."),
            ("Reviewer notes", "Create and resolve field-level notes stored beneath the application."),
            ("Resource management", "Manage matter files, notes, links, global legacy resources, and visa-specific resource templates."),
            ("Document review", "Upload and preview a PDF, then reconcile issues from Firestore and Zoho CRM Corrections."),
        ],
    )

    document.add_heading("1.2 System boundary", level=2)
    add_callout(
        document,
        "Architecture",
        "Browser UI -> Next.js App Router pages and route handlers -> Firestore for structured records and metadata; Zoho CRM for Deal fields and Corrections; Zoho WorkDrive for uploaded file bytes and share links.",
    )
    add_text(
        document,
        "The browser does not directly query Firestore. Pages call local /api route handlers. Those handlers use the Firebase Admin SDK, which operates with service-account privileges and bypasses Firestore security rules. Zoho calls are also server-side through a shared Axios client.",
    )

    document.add_heading("1.3 Runtime components", level=2)
    components = [
        ("Presentation", "Client components under src/app and src/components render the landing page, matter workspace, questionnaire, resources, and admin dashboards."),
        ("Application API", "Next.js route handlers validate requests, resolve matter IDs, serialize Firestore timestamps, and coordinate external writes."),
        ("Domain helpers", "src/lib contains matter resolution, questionnaire structure/progress, resource normalization, session helpers, and the Zoho client."),
        ("Persistent records", "Firestore stores applications, questionnaires, completion maps, comments, notifications, resource metadata, and template metadata."),
        ("External content", "Zoho WorkDrive stores uploaded files; Zoho CRM stores Deal data, final-document URLs, and Corrections records."),
    ]
    add_table_caption(document, "Logical architecture")
    add_table(document, ["Layer", "Responsibility"], components, [1980, 7380])

    document.add_heading("1.4 Matter identity and duplicate handling", level=2)
    add_numbered(
        document,
        [
            "Attempt a direct Firestore lookup at applications/{matterId}.",
            "If no direct document exists, query zohoId, then zohoDealId, then dealId for an exact match.",
            "If more than one application matches, hydrate questionnaire and completion data for each candidate.",
            "Choose the best candidate by most recent timestamp, then completed-key count, profile count, data-document presence, and finally application ID.",
            "Return the canonical Firebase application ID so the matter layout can replace a Zoho-ID URL with a Firebase-ID URL.",
        ],
    )

    # 2. Features
    document.add_heading("2. Feature catalog and user journeys", level=1)
    add_section_intro(
        document,
        "This section describes behavior visible in the current source, including differences between the generic questionnaire review and the subclass 482/186 review.",
    )
    document.add_heading("2.1 Landing page and matter workspace", level=2)
    add_bullets(
        document,
        [
            "Search by Firebase Application ID or Zoho Deal ID.",
            "Redirect directly to the matter questionnaire.",
            "Show reference, normalized visa label, Firebase ID, Zoho Deal ID, and completion percentage.",
            "Use a sticky responsive header with Questionnaire, Resources, and Document Review tabs.",
            "Collapse the matter summary after scrolling while keeping navigation visible.",
            "Display database and not-found errors with a return-to-search action.",
        ],
    )

    document.add_heading("2.2 Questionnaire review", level=2)
    add_text(
        document,
        "Questionnaire data is read from applications/{appId}/data/questionnaire. The UI selects one of two review experiences:",
    )
    questionnaire_rows = [
        (
            "Skills in Demand / Employer Nomination",
            "Triggered by temporary-work, visaContext 482/186, or matching labels. Provides profile-aware sidebar navigation, completion markers, read-only field controls, Previous/Next movement, included-applicant summaries, and all-applicant sections.",
        ),
        (
            "Generic structured review",
            "Used for other matters. Builds applicant, all-applicant, non-migrating, and uncategorized sections. Provides profile/section tabs, field-level reviewer notes, and browser print export with an option to include notes.",
        ),
    ]
    add_table_caption(document, "Questionnaire review modes")
    add_table(document, ["Mode", "Implemented behavior"], questionnaire_rows, [2880, 6480])
    add_bullets(
        document,
        [
            ("Main applicant topics", "Details, other names, identity, contact details, employment, education, skills, and language."),
            ("Spouse topics", "Details, other names, and identity; subclass 186 also adds education and language."),
            ("Child topics", "Details, other names, identity, and custody."),
            ("All-applicant topics", "Visas, travel history, countries of residence, health, and character."),
            ("Progress", "Temporary-work progress is calculated against profile-aware expected sections; other visa types use the configured route catalog."),
        ],
    )

    document.add_heading("2.3 Reviewer comments", level=2)
    add_bullets(
        document,
        [
            "Create a comment against a field path with label, body, severity, source, and optional document URL.",
            "Store open/resolved status, a derived section key, author placeholders, and timestamps.",
            "List comments in creation order and optionally filter by source.",
            "Edit body or severity and change status between open and resolved.",
            "Delete Firestore comments.",
            "When source=documentReview, merge Zoho CRM Corrections into the read response and update their status in Zoho when changed.",
        ],
    )

    document.add_heading("2.4 Matter-specific resources", level=2)
    add_bullets(
        document,
        [
            "Create files, links, or notes for one application.",
            "Search the active matter resource list by title, description, note content, file name, or URL.",
            "Upload files up to 50 MB to the WorkDrive folder referenced by the Zoho Deal's Workdrive_Folder_ID or legacy WorkDrive_Folder_ID field.",
            "Store resource metadata under the application in Firestore.",
            "Archive a resource; if it has a WorkDrive resource ID, the external resource is moved to WorkDrive deleted status before Firestore is marked archived.",
            "Embed, Library, and Folder buttons are shown as disabled future options on the individual-resource screen.",
        ],
    )

    document.add_heading("2.5 Visa-specific shared resource templates", level=2)
    add_text(
        document,
        "The All Matters tab embeds the template manager. Templates are seeded on demand for four visa slugs and use fixed WorkDrive root folders with environment-variable overrides.",
    )
    template_rows = [
        ("partner", "Subclass 820/801", "Partner template root"),
        ("protection", "Subclass 866", "Protection template root"),
        ("482", "Subclass 482", "Skills in Demand template root"),
        ("186", "Subclass 186", "Employer Nomination template root"),
    ]
    add_table_caption(document, "Resource templates")
    add_table(document, ["Visa slug", "Display title", "WorkDrive ownership"], template_rows, [1440, 2880, 5040])
    add_bullets(
        document,
        [
            "View all templates or filter to one visa type.",
            "Create named categories with an icon; categories are stored on the template document.",
            "Create file, link, note, and logical folder items.",
            "Upload multiple files and retain the public link returned by WorkDrive.",
            "Search, sort, edit metadata, hide/reactivate, move between logical parents, and delete items.",
            "Reject missing/non-folder parents and prevent folder cycles.",
            "Delete WorkDrive content before deleting a file item from Firestore.",
        ],
    )
    add_callout(
        document,
        "Important",
        "Template folder items and categories are logical Firestore organization. File uploads go to the configured visa root folder; the current implementation does not create corresponding WorkDrive subfolders.",
        "warning",
    )

    document.add_heading("2.6 Document review", level=2)
    add_bullets(
        document,
        [
            "Accept only PDF uploads for the document-review source, verifying both the .pdf extension and the %PDF- file signature.",
            "Enforce a 50 MB limit.",
            "Find or create a matter-named folder beneath the configured document-review WorkDrive root.",
            "Upload the file, create a public WorkDrive link, derive an embed URL, and preview it in an iframe.",
            "Store the resource under the application and copy Final_File_For_Visa_Submission to both Firestore and the Zoho Deal.",
            "Show open and resolved issues from Firestore plus Zoho Corrections.",
            "Archive the current review resource and clear the final-file field in Firestore and Zoho.",
        ],
    )

    document.add_heading("2.7 Admin dashboards", level=2)
    admin_rows = [
        ("/admin/document-review", "Aggregate Firestore document-review comments across applications; filter/search and resolve/reopen."),
        ("/admin/resources", "Legacy global resource library with file/link/note CRUD, statuses, scopes, metadata filters, and quick publishing actions."),
        ("/login", "Admin-key form exists, but current getAdminSession behavior redirects immediately because it always returns an admin actor."),
    ]
    add_table_caption(document, "Admin-facing features")
    add_table(document, ["Page", "Purpose"], admin_rows, [2520, 6840])

    # 3. Pages and stack
    document.add_heading("3. Application pages and technical stack", level=1)
    document.add_heading("3.1 Page routes", level=2)
    page_rows = [
        ("/", "Matter lookup", "Enter a Firebase application ID or Zoho Deal ID."),
        ("/matter/[matterId]", "Redirect", "Redirects to the questionnaire."),
        ("/matter/[matterId]/questionnaire", "Questionnaire review", "Read answers, progress, and supported reviewer notes/export."),
        ("/matter/[matterId]/resources", "Resources", "Manage visa templates or resources limited to one matter."),
        ("/matter/[matterId]/document-review", "Document review", "Upload/preview PDF and manage issue status."),
        ("/admin", "Redirect", "Redirects to the legacy resource library."),
        ("/admin/resources", "Legacy library", "Global resource management."),
        ("/admin/document-review", "Corrections dashboard", "Cross-matter Firestore correction list."),
        ("/login", "Admin login", "Admin-key form; currently bypassed by the session helper."),
    ]
    add_table_caption(document, "User-interface route map")
    add_table(document, ["Route", "Screen", "Responsibility"], page_rows, [3060, 1800, 4500])

    document.add_heading("3.2 Core technology", level=2)
    stack_rows = [
        ("Framework", "Next.js 16.2.4", "App Router pages and route handlers; dynamic params are awaited."),
        ("UI runtime", "React 19.2.4 / React DOM 19.2.4", "Client-side state, effects, memoization, and responsive views."),
        ("Language", "JavaScript, JSX, and ESM", "No TypeScript source; @/* maps to ./src/*."),
        ("Styling", "Tailwind CSS 4 + @tailwindcss/postcss", "Utility-first layout and the project color system."),
        ("UI primitives", "Radix UI", "Collapsible, Progress, Scroll Area, Slot, and related local wrappers."),
        ("Icons", "lucide-react 1.11", "Interface icons."),
        ("Class utilities", "clsx, tailwind-merge, class-variance-authority", "Conditional classes and component variants."),
        ("Database SDK", "firebase-admin 13.8", "Server-side Firestore and Auth service initialization."),
        ("HTTP client", "Axios 1.15.2", "Zoho token, CRM, and WorkDrive requests."),
        ("Multipart", "form-data 4.0.5", "Server-side WorkDrive and CRM attachment uploads."),
        ("Quality", "ESLint 9 + eslint-config-next 16.2.4", "Next core-web-vitals lint configuration."),
        ("Hosting metadata", "Firebase configuration present", "Firestore region asia-east2, rules, and indexes are tracked."),
    ]
    add_table_caption(document, "Technology stack")
    add_table(document, ["Area", "Technology", "Role"], stack_rows, [1800, 2520, 5040])

    document.add_heading("3.3 Installed but not active in the inspected application flow", level=2)
    add_bullets(
        document,
        [
            "The Firebase client package is installed, but application data access in src uses firebase-admin on the server.",
            "date-fns and @radix-ui/react-tabs are declared dependencies but are not imported by current src files.",
            "DATABASE_URL and CUSTOM_DATABASE_URL are present in local environment files but are not referenced by application source.",
        ],
    )

    # 4. Internal APIs
    document.add_heading("4. Internal API reference", level=1)
    add_section_intro(
        document,
        "All endpoints are Next.js route handlers. JSON responses use a success boolean and usually include an error message on failure.",
    )
    api_rows = [
        ("POST", "/api/admin/session", "Validate PORTAL_ADMIN_KEY and set a signed 12-hour HTTP-only cookie."),
        ("DELETE", "/api/admin/session", "Clear the admin session cookie."),
        ("GET", "/api/admin/document-review", "Collection-group query for Firestore reviewComments where source=documentReview."),
        ("GET", "/api/matter/[matterId]", "Resolve the matter, read application/questionnaire/completion, and calculate progress."),
        ("GET", "/api/matter/[matterId]/resources", "List non-archived resources beneath the resolved application."),
        ("POST", "/api/matter/[matterId]/resources", "Create a file, link, or note; files go to WorkDrive."),
        ("PATCH", "/api/matter/[matterId]/resources/[resourceId]", "Archive a resource; delete external WorkDrive content when an ID is present."),
        ("GET", "/api/resources", "List/filter/search the legacy global resource collection."),
        ("POST", "/api/resources", "Create a global file, link, or note."),
        ("PATCH", "/api/resources/[resourceId]", "Edit global resource metadata, URL/note content, scope, or status."),
        ("DELETE", "/api/resources/[resourceId]", "Delete the Firestore global resource document."),
        ("GET", "/api/resource-templates", "Ensure the four templates exist and return item counts."),
        ("GET", "/api/resource-templates/[visaSlug]", "Return one template and its sorted items."),
        ("PATCH", "/api/resource-templates/[visaSlug]", "Update title, status, or categories."),
        ("POST", "/api/resource-templates/[visaSlug]/items", "Create folder/file/link/note item and update the template timestamp."),
        ("PATCH", "/api/resource-templates/[visaSlug]/items/[itemId]", "Edit name, parent, category, order, status, link, or note."),
        ("DELETE", "/api/resource-templates/[visaSlug]/items/[itemId]", "Delete WorkDrive resource when present, then delete Firestore item."),
        ("GET", "/api/review-comments/[matterId]", "List Firestore comments; source=documentReview also merges Zoho Corrections."),
        ("POST", "/api/review-comments/[matterId]", "Create a Firestore review comment."),
        ("PATCH", "/api/review-comments/[matterId]/[commentId]", "Update a Firestore comment or Zoho Correction status."),
        ("DELETE", "/api/review-comments/[matterId]/[commentId]", "Delete a Firestore review comment."),
    ]
    add_table_caption(document, "Complete internal endpoint catalog")
    add_table(document, ["Method", "Route", "Purpose"], api_rows, [900, 3600, 4860])

    document.add_heading("4.1 Request and validation rules", level=2)
    validation_rows = [
        ("Matter ID", "Required; Firebase ID is tried first, followed by exact Zoho-ID fields."),
        ("Matter files", "Maximum 50 MB; arbitrary type except document-review uploads."),
        ("Document-review file", "Maximum 50 MB; .pdf extension and PDF header are both required."),
        ("Links", "Only http and https URLs are accepted."),
        ("Notes", "Non-empty note text is required."),
        ("Template visa slugs", "partner, protection, 482, and 186, with limited aliases normalized."),
        ("Template parent", "Must be a folder in the same template; folder moves cannot form a cycle."),
        ("Comment status", "Only open or resolved."),
        ("Resource archive", "Matter-specific resource PATCH supports only status=archived."),
    ]
    add_table_caption(document, "Boundary validation")
    add_table(document, ["Input", "Rule"], validation_rows, [2520, 6840])

    document.add_heading("4.2 Response and error behavior", level=2)
    add_bullets(
        document,
        [
            "400 for missing or invalid request values.",
            "401 only where explicit session/key checks run; current session helper normally returns an admin actor.",
            "404 for unresolved matters, resources, templates, or items.",
            "500 for initialization and unexpected server errors.",
            "502 where Zoho/WorkDrive cannot supply required remote data or a usable public link.",
            "Firestore timestamps are converted to ISO strings before responses are serialized.",
        ],
    )

    # 5. External integrations
    document.add_heading("5. External APIs and integrations", level=1)
    document.add_heading("5.1 Firebase Admin SDK", level=2)
    add_text(
        document,
        "src/lib/firebase-admin.js initializes a singleton Firebase app on module load. Credential discovery is ordered for hosting environments first and local files last.",
    )
    add_numbered(
        document,
        [
            "Decode FIREBASE_SERVICE_ACCOUNT_BASE64, if present.",
            "Parse FIREBASE_SERVICE_ACCOUNT_KEY as raw JSON.",
            "Load GOOGLE_APPLICATION_CREDENTIALS or FIREBASE_SERVICE_ACCOUNT_PATH.",
            "Fallback to service.json or service-account.json in the repository root.",
            "Use NEXT_PUBLIC_FIREBASE_PROJECT_ID or project_id from the service account.",
        ],
    )
    add_callout(
        document,
        "Security",
        "The service account grants server-side administrative access. The service JSON is ignored by Git and must not be distributed with application documentation or client bundles.",
        "warning",
    )

    document.add_heading("5.2 Zoho access-token broker", level=2)
    add_text(
        document,
        "The application does not perform an OAuth refresh-token exchange itself. It performs an HTTP GET to ACCESSTOKEN_URL or ZOHO_ACCESS_TOKEN_URL; WorkDrive can use WORKDRIVE_ACCESSTOKEN_URL first. The broker may return access_token, details.output, output, or a plain token string.",
    )
    add_bullets(
        document,
        [
            "The Zoho-oauthtoken prefix is stripped when received and added to Authorization headers on outbound calls.",
            "CRM and WorkDrive tokens are cached independently in memory for 50 minutes.",
            "CRM requests retry once with a refreshed token after HTTP 401.",
            "The datacenter is controlled by ZOHO_DATACENTER and defaults to com.au.",
        ],
    )

    document.add_heading("5.3 Zoho CRM APIs", level=2)
    crm_rows = [
        ("CRM v7 base", "https://www.zohoapis.{datacenter}/crm/v7", "Record search/CRUD, COQL, and related lists."),
        ("Deals", "GET /Deals/{id}", "Read WorkDrive folder fields and general Deal data."),
        ("Deals update", "PUT /Deals", "Set or clear Final_File_For_Visa_Submission."),
        ("Corrections", "GET/PUT /Deals/{id}/Corrections", "Read document-review issues and update open/resolved status."),
        ("CRM v8 attachments", "POST /crm/v8/{module}/{id}/Attachments", "Generic attachment helper available in the client wrapper."),
    ]
    add_table_caption(document, "Zoho CRM endpoints used or exposed by the shared client")
    add_table(document, ["Capability", "Endpoint pattern", "Portal responsibility"], crm_rows, [2160, 3240, 3960])
    add_text(
        document,
        "The shared Zoho client also contains generic contact search, record creation/update, related-record CRUD, COQL, and Partner_Dependents synchronization helpers. Those helpers are available infrastructure but are not called by the current portal route set except for Deals and Corrections operations.",
    )

    document.add_heading("5.4 Zoho WorkDrive API", level=2)
    workdrive_rows = [
        ("List children", "GET /files/{folderId}/files", "Find an existing matter folder; scans in 50-item pages up to 1,000 children."),
        ("Create folder", "POST /files", "Create the document-review matter folder."),
        ("Upload file", "POST /upload", "Store file bytes in a configured folder."),
        ("Create public link", "POST /links", "Create an external link with download allowed."),
        ("Delete resource", "PATCH /files/{resourceId}", "Set WorkDrive status 51; 404 is treated as already missing."),
    ]
    add_table_caption(document, "WorkDrive operations")
    add_table(document, ["Operation", "Endpoint", "Use"], workdrive_rows, [2160, 2880, 4320])
    add_callout(
        document,
        "Data exposure",
        "Uploaded resource URLs are public/external WorkDrive links. Treat the returned URL as shareable access and apply the organization's WorkDrive policy accordingly.",
        "warning",
    )

    # 6. Storage
    document.add_heading("6. Data storage and information ownership", level=1)
    add_section_intro(
        document,
        "Structured data and file content are intentionally split. Firestore is the system of record for portal metadata; WorkDrive owns file bytes; Zoho CRM owns Deal and Corrections records.",
    )
    document.add_heading("6.1 Firestore collections used by application code", level=2)
    firestore_rows = [
        ("applications/{appId}", "Core matter metadata", "Reference, applicant/client fields, visa type, Zoho IDs, timestamps, and final document URL."),
        ("applications/{appId}/data/questionnaire", "Questionnaire answers", "Profiles, profile data, all-applicant answers, visaContext, and legacy keys."),
        ("applications/{appId}/data/completion", "Progress map", "Boolean completion keys and timestamps."),
        ("applications/{appId}/reviewComments/{commentId}", "Review comments", "Path, label, body, severity, source, status, section key, author placeholders, and timestamps."),
        ("applications/{appId}/resources/{resourceId}", "Matter resources", "File/link/note metadata, URLs, WorkDrive IDs, status, source, category, and timestamps."),
        ("notifications/{notificationId}", "Applicant notification intent", "Application ID, type, title/body, path, read flag, and createdAt."),
        ("resources/{resourceId}", "Legacy global library", "Type, title/content, scope, status, category/program/audience, WorkDrive metadata, and lifecycle timestamps."),
        ("resourceTemplates/{visaSlug}", "Template header", "Title, status, categories, WorkDrive root, actor metadata, and timestamps."),
        ("resourceTemplates/{visaSlug}/items/{itemId}", "Template items", "Logical parent, kind, name, category, order, visibility, content/URL, WorkDrive metadata, and timestamps."),
    ]
    add_table_caption(document, "Application-owned Firestore paths")
    add_table(document, ["Path", "Purpose", "Representative content"], firestore_rows, [3420, 1980, 3960])

    document.add_heading("6.2 Collections present in Firestore rules but not used by current portal source", level=2)
    add_bullets(
        document,
        [
            "users/{userId} and nested preferences, drafts, and zoho documents.",
            "messages/{messageId}.",
            "Authenticated-owner and publicReviewAccess rules for applications and their subcollections.",
        ],
    )
    add_text(
        document,
        "These rules appear to support an adjacent applicant/client portal. The admin portal uses Firebase Admin on the server, so these client security rules do not restrict its route handlers.",
    )

    document.add_heading("6.3 Zoho CRM ownership", level=2)
    zoho_storage_rows = [
        ("Deals record", "Matter/CRM identity and fields used for WorkDrive folder lookup."),
        ("Workdrive_Folder_ID / WorkDrive_Folder_ID", "Folder ID for normal matter-specific file uploads."),
        ("Final_File_For_Visa_Submission", "Preview/embed URL for the current review document; updated in both Zoho and Firestore."),
        ("Corrections related list", "Client-submitted document issues, including field/name, issue description, status, email, and timestamps."),
    ]
    add_table_caption(document, "Zoho CRM data used by the portal")
    add_table(document, ["Zoho location", "Stored information"], zoho_storage_rows, [3240, 6120])

    document.add_heading("6.4 WorkDrive ownership", level=2)
    add_bullets(
        document,
        [
            "Actual bytes for matter-specific files.",
            "Actual bytes for global legacy-library files.",
            "Actual bytes for visa-template files.",
            "Document-review PDFs in matter-named folders beneath the review root.",
            "External link objects and downloadable/public URLs.",
        ],
    )

    document.add_heading("6.5 Browser and process storage", level=2)
    browser_rows = [
        ("Cookie", "vp_admin_session", "Signed HTTP-only, SameSite=Lax, 12-hour maximum age; secure in production."),
        ("React state", "Ephemeral UI state", "Filters, selected resources, form values, open drawers, and current navigation."),
        ("Server memory", "Zoho token cache", "Separate CRM and WorkDrive access tokens for up to 50 minutes."),
        ("Environment/files", "Credentials and configuration", "Server-only environment variables or ignored service-account JSON files."),
    ]
    add_table_caption(document, "Non-database storage")
    add_table(document, ["Location", "Key/data", "Purpose"], browser_rows, [1800, 2520, 5040])

    # 7. Data flows
    document.add_heading("7. Core data flows", level=1)
    document.add_heading("7.1 Open a matter", level=2)
    add_numbered(
        document,
        [
            "The landing page navigates to /matter/{input}/questionnaire.",
            "The matter layout calls GET /api/matter/{input}.",
            "The route resolves the Firebase application, reads questionnaire and completion documents, and calculates progress.",
            "The layout replaces a Zoho-ID URL with the canonical Firebase ID when needed.",
            "Child pages call their own APIs for comments, resources, or document-review data.",
        ],
    )

    document.add_heading("7.2 Upload a normal matter resource", level=2)
    add_numbered(
        document,
        [
            "Resolve the application and derive its Zoho Deal ID.",
            "Fetch the Deal and read Workdrive_Folder_ID or WorkDrive_Folder_ID.",
            "Upload bytes to WorkDrive and create a public link.",
            "Store metadata and WorkDrive identifiers beneath applications/{appId}/resources.",
            "Return the serialized resource to the browser.",
        ],
    )

    document.add_heading("7.3 Upload a document-review PDF", level=2)
    add_numbered(
        document,
        [
            "Verify the 50 MB limit, .pdf extension, and PDF signature.",
            "Resolve the configured review root and find or create a matter-named WorkDrive folder.",
            "Upload the PDF and create a public link plus embed URL.",
            "Update the Zoho Deal's Final_File_For_Visa_Submission field.",
            "Commit the Firestore resource and application final-file field in one batch.",
            "Preview the document and merge Firestore document-review comments with Zoho Corrections.",
        ],
    )

    document.add_heading("7.4 Publish a visa-template file", level=2)
    add_numbered(
        document,
        [
            "Ensure resourceTemplates/{visaSlug} exists and has the current root-folder metadata.",
            "Validate kind, status, category, order, and logical parent.",
            "Upload file bytes directly to the configured visa WorkDrive root.",
            "Create a public link and store the item under resourceTemplates/{visaSlug}/items.",
            "Update the template's timestamp and actor metadata.",
        ],
    )

    document.add_heading("7.5 Resolve an issue", level=2)
    add_bullets(
        document,
        [
            ("Firestore comment", "PATCH updates status and updatedAt on the comment document."),
            ("Zoho correction", "IDs prefixed zohoCorrection: are translated to a Zoho related-record update with Status=Resolved or Open."),
        ],
    )

    document.add_heading("7.6 Deletion and archive semantics", level=2)
    delete_rows = [
        ("Matter resource", "External WorkDrive resource is deleted when possible; Firestore document remains with status=archived."),
        ("Document-review resource", "Same as matter resource, plus final-file fields are cleared in Firestore and Zoho."),
        ("Template item", "WorkDrive resource is deleted first, then the Firestore item is permanently deleted."),
        ("Legacy global resource", "Firestore document is permanently deleted; the route does not delete the corresponding WorkDrive file."),
        ("Review comment", "Firestore comment is permanently deleted; Zoho-prefixed corrections are not handled by DELETE."),
    ]
    add_table_caption(document, "Lifecycle behavior")
    add_table(document, ["Object", "Current behavior"], delete_rows, [2880, 6480])

    # 8. Auth/security
    document.add_heading("8. Authentication, authorization, and security", level=1)
    document.add_heading("8.1 Implemented session mechanism", level=2)
    add_text(
        document,
        "The repository includes an admin-key login mechanism. A valid PORTAL_ADMIN_KEY is compared with timingSafeEqual, then a role=admin payload is signed with HMAC-SHA256 using SESSION_SECRET. The cookie is HTTP-only, SameSite=Lax, secure in production, and expires after 12 hours.",
    )
    add_callout(
        document,
        "Current behavior",
        "getAdminSession() does not read or verify the cookie. It always returns a fresh admin actor, and requireAdminSession() delegates to it. Therefore /login redirects immediately and routes that call these helpers are effectively open.",
        "risk",
    )

    document.add_heading("8.2 Route access posture", level=2)
    access_rows = [
        ("Matter pages and /api/matter/*", "No authentication check."),
        ("/api/review-comments/*", "No authentication check."),
        ("/api/admin/document-review", "Calls requireAdminSession, which currently auto-grants."),
        ("/api/resources*", "Uses getAdminSession only to label the actor; it currently auto-grants."),
        ("/api/resource-templates*", "Checks getAdminSession, which currently auto-grants."),
        ("/api/admin/session", "POST validates the admin key, but the rest of the app does not require the resulting cookie."),
    ]
    add_table_caption(document, "Current authorization posture")
    add_table(document, ["Area", "Enforcement"], access_rows, [3240, 6120])

    document.add_heading("8.3 Security controls that are present", level=2)
    add_bullets(
        document,
        [
            "Server-only Firebase and Zoho credentials.",
            "HMAC-signed session token implementation and timing-safe key comparison.",
            "HTTP/HTTPS URL allow-list for submitted links.",
            "File-name sanitization for WorkDrive uploads.",
            "50 MB upload limits.",
            "Document-review PDF signature verification.",
            "Template parent validation and folder-cycle prevention.",
            "No service-account or .env files tracked by the repository ignore rules.",
        ],
    )

    document.add_heading("8.4 Security actions before external production exposure", level=2)
    add_bullets(
        document,
        [
            "Restore real cookie verification in getAdminSession and requireAdminSession.",
            "Apply the gate consistently to matter reads and every mutation endpoint, or implement a separate authenticated user role model.",
            "Add CSRF protection or an equivalent same-origin mutation strategy once sessions are enforced.",
            "Review whether public WorkDrive links are appropriate for each resource category.",
            "Avoid logging token-response bodies or record samples in production.",
            "Confirm client-portal Firestore rules for resources, resourceTemplates, notifications, and reviewComments if clients will access Firestore directly.",
        ],
    )

    # 9. Env
    document.add_heading("9. Configuration and environment variables", level=1)
    add_section_intro(
        document,
        "Values are intentionally omitted. Configure these in the deployment environment; do not place secrets in NEXT_PUBLIC variables.",
    )
    env_rows = [
        ("NEXT_PUBLIC_FIREBASE_PROJECT_ID", "Firebase", "Required unless project_id is supplied by service account."),
        ("FIREBASE_SERVICE_ACCOUNT_BASE64", "Firebase", "Preferred hosted credential option; base64 JSON."),
        ("FIREBASE_SERVICE_ACCOUNT_KEY", "Firebase", "Raw service-account JSON alternative."),
        ("GOOGLE_APPLICATION_CREDENTIALS", "Firebase", "Path to service-account file."),
        ("FIREBASE_SERVICE_ACCOUNT_PATH", "Firebase", "Application-specific path alternative."),
        ("FIREBASE_PROJECT_ID", "Migration script", "Project ID fallback used by scripts/migrate-shared-resources.js."),
        ("ACCESSTOKEN_URL", "Zoho", "Primary token-broker URL."),
        ("ZOHO_ACCESS_TOKEN_URL", "Zoho", "Backward-compatible token-broker URL."),
        ("WORKDRIVE_ACCESSTOKEN_URL", "Zoho", "Optional WorkDrive-specific token broker."),
        ("ZOHO_DATACENTER", "Zoho", "Optional; defaults to com.au."),
        ("PORTAL_ADMIN_KEY", "Session", "Admin-key endpoint secret."),
        ("SESSION_SECRET", "Session", "HMAC secret for signed cookie."),
        ("DOCUMENT_REVIEW_WORKDRIVE_FOLDER_ID", "WorkDrive", "Document-review root override."),
        ("WORKDRIVE_DOCUMENT_REVIEW_FOLDER_ID", "WorkDrive", "Legacy/alternate review root override."),
        ("SHARED_RESOURCES_WORKDRIVE_FOLDER_ID", "WorkDrive", "Preferred global legacy-library root override."),
        ("WORKDRIVE_SHARED_FOLDER_ID", "WorkDrive", "Global root fallback."),
        ("SHARED_WORKDRIVE_FOLDER_ID", "WorkDrive", "Global root fallback."),
        ("WORKDRIVE_FOLDER_ID", "WorkDrive", "Broad global root fallback."),
        ("RESOURCE_TEMPLATE_PARTNER_WORKDRIVE_FOLDER_ID", "Templates", "Partner template root override."),
        ("PARTNER_RESOURCE_TEMPLATE_WORKDRIVE_FOLDER_ID", "Templates", "Partner alternate override."),
        ("RESOURCE_TEMPLATE_PROTECTION_WORKDRIVE_FOLDER_ID", "Templates", "Protection template root override."),
        ("PROTECTION_RESOURCE_TEMPLATE_WORKDRIVE_FOLDER_ID", "Templates", "Protection alternate override."),
        ("RESOURCE_TEMPLATE_482_WORKDRIVE_FOLDER_ID", "Templates", "Subclass 482 root override."),
        ("VISA_482_RESOURCE_TEMPLATE_WORKDRIVE_FOLDER_ID", "Templates", "Subclass 482 alternate override."),
        ("RESOURCE_TEMPLATE_186_WORKDRIVE_FOLDER_ID", "Templates", "Subclass 186 root override."),
        ("VISA_186_RESOURCE_TEMPLATE_WORKDRIVE_FOLDER_ID", "Templates", "Subclass 186 alternate override."),
    ]
    add_table_caption(document, "Environment variables referenced by source")
    add_table(document, ["Variable", "Area", "Purpose"], env_rows, [3600, 1440, 4320])

    document.add_heading("9.1 Present locally but not used by current application source", level=2)
    unused_env_rows = [
        ("NEXT_PUBLIC_FIREBASE_API_KEY", "Configured locally; no src reference."),
        ("NEXT_PUBLIC_FIREBASE_APP_ID", "Configured locally; no src reference."),
        ("NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN", "Configured locally; no src reference."),
        ("NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID", "Configured locally; no src reference."),
        ("NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET", "Configured locally; no src reference."),
        ("DATABASE_URL", "Configured locally; no src reference."),
        ("CUSTOM_DATABASE_URL", "Configured locally; no src reference."),
    ]
    add_table_caption(document, "Configured-but-unreferenced variable names")
    add_table(document, ["Variable", "Status"], unused_env_rows, [3960, 5400])

    # 10. Folder structure
    document.add_heading("10. Folder structure and source ownership", level=1)
    add_section_intro(
        document,
        "Generated folders such as node_modules and .next are omitted. The tree below focuses on maintained source and operational configuration.",
    )
    tree = r"""validify-pro-admin-portal/
|-- public/                         Static logos, favicon, and starter SVG assets
|-- scripts/
|   `-- migrate-shared-resources.js  Matter-resource to global-library migration
|-- src/
|   |-- app/
|   |   |-- admin/                  Admin layouts and dashboard pages
|   |   |-- api/                    Next.js route handlers
|   |   |-- login/                  Admin-key form route
|   |   |-- matter/[matterId]/      Matter shell, questionnaire, resources, review
|   |   |-- globals.css             Tailwind import and print/responsive styling
|   |   |-- layout.js               Root metadata and HTML shell
|   |   `-- page.js                 Matter lookup landing page
|   |-- components/
|   |   |-- admin/                  Admin managers and session controls
|   |   |-- ui/                     Local reusable UI primitives
|   |   |-- QuestionnaireSidebar.jsx
|   |   `-- SkillsInDemandQuestionnaireReview.jsx
|   `-- lib/                        Firebase, Zoho, domain rules, and helpers
|-- .env / .env.local               Local configuration; ignored by Git
|-- firebase.json                   Firestore deployment configuration
|-- firestore.rules                 Client Firestore security rules
|-- firestore.indexes.json          Messages indexes
|-- next.config.mjs                 Next.js configuration
|-- package.json                    Dependencies and npm scripts
|-- postcss.config.mjs              Tailwind PostCSS plugin
`-- jsconfig.json                   @/* -> src/* import alias"""
    add_code_block(document, tree)

    document.add_heading("10.1 API route ownership", level=2)
    api_tree = r"""src/app/api/
|-- admin/
|   |-- document-review/route.js     Cross-application Firestore corrections
|   `-- session/route.js             Admin-key cookie create/delete
|-- matter/[matterId]/
|   |-- route.js                     Matter + questionnaire + progress
|   `-- resources/
|       |-- route.js                 List/create matter resources
|       `-- [resourceId]/route.js    Archive matter resource
|-- resource-templates/
|   |-- route.js                     Seed/list templates
|   `-- [visaSlug]/
|       |-- route.js                 Get/update template
|       `-- items/
|           |-- route.js             Create template item
|           `-- [itemId]/route.js    Update/delete item
|-- resources/
|   |-- route.js                     Legacy global list/create
|   `-- [resourceId]/route.js        Legacy update/delete
`-- review-comments/[matterId]/
    |-- route.js                     List/create comments + Zoho merge
    `-- [commentId]/route.js         Update/delete comments"""
    add_code_block(document, api_tree)

    document.add_heading("10.2 Key libraries", level=2)
    lib_rows = [
        ("adminSession.js", "Admin-key verification, HMAC token creation, cookie options, and current auto-grant helpers."),
        ("firebase-admin.js", "Credential discovery and singleton Firebase Admin/Firestore initialization."),
        ("matterResolver.js", "Firebase/Zoho ID resolution and duplicate selection."),
        ("pdfUploadRules.mjs", "PDF header and extension validation."),
        ("questionnaireProgress.js", "Profile-aware temporary-work progress calculation."),
        ("questionnaireSections.js", "Convert raw questionnaire keys into applicant/shared/non-migrating sections."),
        ("resourceTemplates.js", "Template definitions, normalization, Firestore seed/read helpers, WorkDrive upload/delete."),
        ("routes.js", "Partner, protection, subclass 482, and subclass 186 intake route catalogs."),
        ("sharedResources.js", "Global resource normalization, serialization, folder selection, and WorkDrive upload."),
        ("visaDisplay.js", "Visa-label normalization and Zoho deal type mapping."),
        ("workDrivePreviewUrl.mjs", "Convert supported public WorkDrive links to embed URLs."),
        ("zohoClient.js", "Zoho token, CRM, related-list, WorkDrive, and attachment client."),
    ]
    add_table_caption(document, "Library responsibilities")
    add_table(document, ["File", "Responsibility"], lib_rows, [3060, 6300])

    document.add_heading("10.3 Local UI component layer", level=2)
    add_text(
        document,
        "src/components/ui contains local wrappers for alerts, badges, buttons, cards, checkboxes, collapsibles, dialogs, inputs, labels, progress, radio groups, scroll areas, selects, textareas, toasts, and the toaster. Admin managers compose these primitives rather than introducing a separate component framework.",
    )

    # 11. Operations
    document.add_heading("11. Development, deployment, and operations", level=1)
    document.add_heading("11.1 Local commands", level=2)
    commands = r"""npm install
npm run dev
npm run lint
npm run build
npm run start
node --test src/lib/pdfUploadRules.test.mjs src/lib/workDrivePreviewUrl.test.mjs"""
    add_code_block(document, commands)
    command_rows = [
        ("npm run dev", "Start the Next.js development server."),
        ("npm run build", "Create the production Next.js build."),
        ("npm run start", "Run the production server after a build."),
        ("npm run lint", "Run ESLint with Next core-web-vitals rules."),
        ("npm run migrate:shared-resources", "Migrate legacy per-application resources to the global resources collection."),
    ]
    add_table_caption(document, "Package scripts")
    add_table(document, ["Command", "Purpose"], command_rows, [3240, 6120])

    document.add_heading("11.2 Shared-resource migration", level=2)
    add_text(
        document,
        "scripts/migrate-shared-resources.js scans every application resources subcollection and copies eligible records to the global resources collection. It records migration.legacyResourcePath to avoid copying the same source twice.",
    )
    add_bullets(
        document,
        [
            "--dry-run reports work without writing.",
            "--include-archived includes archived source resources; they are skipped by default.",
            "Writes are committed in batches of 200.",
            "The script reads .env and .env.local and supports the same Firebase credential fallbacks.",
        ],
    )

    document.add_heading("11.3 Firestore deployment configuration", level=2)
    add_bullets(
        document,
        [
            "Default database in region asia-east2.",
            "Rules file: firestore.rules.",
            "Indexes file: firestore.indexes.json.",
            "Tracked composite indexes cover messages by userId and createdAt ascending/descending.",
        ],
    )
    add_callout(
        document,
        "Deployment note",
        "The repository does not encode one production hosting provider. Next.js can be deployed to a compatible Node.js platform; Firebase configuration in this repository is specifically for Firestore rules and indexes.",
    )

    document.add_heading("11.4 Operational behavior", level=2)
    operations_rows = [
        ("Zoho unavailable", "CRM helper often returns null/empty data when no token is available; WorkDrive mutations throw because file operations cannot continue."),
        ("Duplicate matter IDs", "Resolver selects one candidate deterministically and returns duplicateCount and duplicateIds."),
        ("Remote cleanup failure", "Some upload flows delete the new WorkDrive resource when a later critical step fails."),
        ("Document replacement", "Uploading a replacement adds a new active resource; it does not automatically archive all prior review resources."),
        ("Token lifetime", "Process-memory cache is lost on restart/serverless cold start; the broker is called again."),
    ]
    add_table_caption(document, "Operational characteristics")
    add_table(document, ["Scenario", "Behavior"], operations_rows, [2880, 6480])

    # 12. Verification and risks
    document.add_heading("12. Verification results, limitations, and risks", level=1)
    document.add_heading("12.1 Repository checks performed for this document", level=2)
    verification_rows = [
        ("ESLint", "Passed with 0 errors and 4 warnings."),
        ("Utility tests", "2 passed: PDF upload rules and WorkDrive embed URL."),
        ("Warnings", "Three next/no-img-element warnings and one anonymous default-export warning."),
        ("Source scope", "All maintained files under src, scripts, root configuration, Firestore rules/indexes, and current package metadata were inspected."),
        ("Working tree", "Pre-existing uncommitted application changes were present and were not modified by documentation generation."),
    ]
    add_table_caption(document, "Validation status on 23 July 2026")
    add_table(document, ["Check", "Result"], verification_rows, [2520, 6840])

    document.add_heading("12.2 Known implementation risks", level=2)
    risks = [
        (
            "Critical - authorization bypass",
            "getAdminSession always returns an admin actor, so login and admin guards do not enforce identity. Matter and comment endpoints also have no gate.",
        ),
        (
            "High - public file links",
            "WorkDrive links are created as external/public links with downloads enabled. Access depends on possession of the URL.",
        ),
        (
            "High - legacy file orphaning",
            "DELETE /api/resources/{id} deletes only Firestore metadata and does not remove the associated WorkDrive resource.",
        ),
        (
            "Medium - notification write defect",
            "The review-comment POST path attempts to use appId when creating a non-document-review notification, but appId is not defined in that scope. The notification error is caught, so comment creation can still succeed without a notification.",
        ),
        (
            "Medium - replacement history",
            "A document-review replacement creates another active resource. Selection favors the newest resource, but older active files remain until explicitly archived."),
        (
            "Medium - production logging",
            "The Zoho client logs token response structures and sample CRM records, which can expose sensitive operational data in server logs."),
        (
            "Medium - rules mismatch for direct clients",
            "Tracked Firestore rules do not define client access for global resources, resourceTemplates, notifications, or reviewComments. Server Admin SDK calls still work."),
        (
            "Low - stale auxiliary documentation",
            "Some existing Markdown guides describe routes that are not present in the current source. Use this source-derived document and the route tree as the current baseline.",
        ),
    ]
    add_table_caption(document, "Prioritized risk register")
    add_table(document, ["Risk", "Evidence and impact"], risks, [2700, 6660])

    document.add_heading("12.3 Intentional current limits", level=2)
    add_bullets(
        document,
        [
            "Individual resource Embed, Library, and Folder actions are disabled.",
            "Visa-template category folders do not create WorkDrive subfolders.",
            "WorkDrive child scanning stops after 1,000 items per folder.",
            "Questionnaire answers are read-only; the portal is a review surface, not the intake editor.",
            "The generic print export uses the browser print dialog rather than a server PDF generator.",
            "The admin corrections dashboard aggregates Firestore comments only; Zoho Corrections are merged on the matter document-review screen.",
        ],
    )

    # Appendix
    document.add_heading("Appendix A. Representative data shapes", level=1)
    document.add_heading("A.1 Matter API response", level=2)
    add_code_block(
        document,
        """{
  "success": true,
  "application": { "id": "firebase-app-id", "...": "matter metadata" },
  "questionnaire": { "...": "stored answers" },
  "completion": { "section-key": true },
  "percentage": 0,
  "progress": { "completedSections": 0, "totalSections": 0 },
  "matchedBy": "firebaseId | zohoId | zohoDealId | dealId",
  "duplicateCount": 1,
  "duplicateIds": ["firebase-app-id"]
}""",
    )
    document.add_heading("A.2 Review comment", level=2)
    add_code_block(
        document,
        """{
  "path": "profile:123.details.family_name",
  "label": "Family name",
  "body": "Please confirm the spelling.",
  "severity": "suggestion",
  "source": "questionnaire | documentReview",
  "status": "open | resolved",
  "sectionKey": "profile:123",
  "authorId": "admin | client",
  "authorName": "Admin | Client",
  "createdAt": "timestamp",
  "updatedAt": "timestamp"
}""",
    )
    document.add_heading("A.3 Template item", level=2)
    add_code_block(
        document,
        """{
  "parentId": null,
  "kind": "folder | file | link | note",
  "name": "Resource name",
  "category": "Guides",
  "order": 0,
  "status": "active | hidden",
  "externalUrl": null,
  "workdriveId": null,
  "mimeType": null,
  "size": null
}""",
    )
    document.add_heading("A.4 Matter resource", level=2)
    add_code_block(
        document,
        """{
  "type": "file | link | note",
  "title": "Resource title",
  "description": "",
  "status": "active | archived",
  "source": "documentReview (optional)",
  "category": "Document Review (optional)",
  "publicUrl": "https://...",
  "workDriveResourceId": "resource-id",
  "workDriveFolderId": "folder-id",
  "createdAt": "timestamp",
  "updatedAt": "timestamp"
}""",
    )

    document.add_heading("Appendix B. Source map", level=1)
    source_rows = [
        ("Features and pages", "src/app/**/page.js, src/app/matter/[matterId]/layout.js, src/components/admin/*.jsx"),
        ("Internal APIs", "src/app/api/**/route.js"),
        ("Matter resolution", "src/lib/matterResolver.js"),
        ("Questionnaire structure/progress", "src/lib/questionnaireSections.js, questionnaireProgress.js, routes.js"),
        ("Firebase", "src/lib/firebase-admin.js, firestore.rules, firestore.indexes.json, firebase.json"),
        ("Zoho CRM/WorkDrive", "src/lib/zohoClient.js, sharedResources.js, resourceTemplates.js"),
        ("Authentication", "src/lib/adminSession.js, src/app/api/admin/session/route.js"),
        ("Dependencies", "package.json, package-lock.json, pnpm-lock.yaml"),
        ("Migration", "scripts/migrate-shared-resources.js"),
        ("Configuration", "next.config.mjs, postcss.config.mjs, jsconfig.json, eslint.config.mjs"),
    ]
    add_table_caption(document, "Primary evidence used")
    add_table(document, ["Documentation area", "Repository source"], source_rows, [2700, 6660])
    add_callout(
        document,
        "Maintenance rule",
        "Update this documentation when route files, Firestore paths, Zoho field names, WorkDrive folder ownership, authentication behavior, or package versions change.",
    )

    # Preserve widows/orphans and heading pagination.
    for paragraph in document.paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        widow = p_pr.find(qn("w:widowControl"))
        if widow is None:
            widow = OxmlElement("w:widowControl")
            p_pr.append(widow)
        widow.set(qn("w:val"), "1")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
